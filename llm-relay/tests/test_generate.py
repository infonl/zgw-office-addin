# SPDX-FileCopyrightText: 2025 INFO.nl
# SPDX-License-Identifier: EUPL-1.2+

import json
from unittest.mock import AsyncMock, patch

import httpx


def test_generate_missing_api_key(client, sample_payload):
    """Without OPENROUTER_API_KEY, returns a clear error."""
    response = client.post("/api/v1/generate", json=sample_payload)
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is False
    assert "OPENROUTER_API_KEY" in data["error"]


def test_generate_validation_missing_content(client):
    """Missing required fields returns 422."""
    response = client.post("/api/v1/generate", json={"prompt": "hello"})
    assert response.status_code == 422


def test_generate_validation_missing_prompt(client, b64_content):
    """Missing prompt returns 422."""
    response = client.post(
        "/api/v1/generate",
        json={"content": b64_content, "output_schema": {"summary": "str"}},
    )
    assert response.status_code == 422


def test_generate_validation_missing_output_schema(client, b64_content):
    """Missing output_schema returns 422."""
    response = client.post(
        "/api/v1/generate",
        json={"content": b64_content, "prompt": "Summarize."},
    )
    assert response.status_code == 422


def test_generate_model_is_optional(client, sample_payload):
    """model field can be omitted — should not cause 422."""
    assert "model" not in sample_payload
    response = client.post("/api/v1/generate", json=sample_payload)
    # Will fail at API key check, but not at validation
    assert response.status_code == 200


def test_generate_content_type_is_optional(client, b64_content):
    """content_type can be omitted."""
    payload = {
        "content": b64_content,
        "prompt": "Summarize.",
        "output_schema": {"summary": "str"},
    }
    response = client.post("/api/v1/generate", json=payload)
    assert response.status_code == 200


def _mock_openrouter_response(data: dict) -> httpx.Response:
    body = {"choices": [{"message": {"content": json.dumps(data)}}]}
    return httpx.Response(200, json=body, request=httpx.Request("POST", "https://fake"))


def test_generate_success(client, sample_payload, mock_settings_attrs):
    """Successful round-trip with mocked OpenRouter."""
    llm_data = {"summary": "About housing.", "topics": ["housing", "policy"]}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response),
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)

        response = client.post("/api/v1/generate", json=sample_payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["data"] == llm_data
    assert data["error"] is None


def test_generate_with_explicit_model(client, sample_payload, mock_settings_attrs):
    """Explicit model in request is used."""
    sample_payload["model"] = "mistral/mistral-small-2603"
    llm_data = {"summary": "Test.", "topics": ["test"]}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response),
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)

        response = client.post("/api/v1/generate", json=sample_payload)

    assert response.status_code == 200
    assert response.json()["success"] is True


def test_generate_plain_text_content(client):
    """Plain text (not base64) is accepted as fallback."""
    payload = {
        "content": "This is plain text, not base64 encoded.",
        "prompt": "Summarize.",
        "output_schema": {"summary": "str"},
    }
    response = client.post("/api/v1/generate", json=payload)
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is False
    assert "OPENROUTER_API_KEY" in data["error"]


def test_generate_with_attachment_type(client, sample_payload):
    """attachment_type is accepted without error."""
    sample_payload["attachment_type"] = "item"
    response = client.post("/api/v1/generate", json=sample_payload)
    assert response.status_code == 200


def test_generate_image_content_type(client, mock_settings_attrs):
    """Image content_type uses vision message format."""
    import base64

    # Tiny 1x1 PNG
    pixel = base64.b64encode(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50).decode()
    payload = {
        "content": pixel,
        "content_type": "image/png",
        "prompt": "Describe this image.",
        "output_schema": {"description": "str"},
    }
    llm_data = {"description": "A tiny image."}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response) as mock_post,
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)

        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    assert response.json()["success"] is True

    # Verify the vision message format was used
    call_kwargs = mock_post.call_args
    sent_payload = call_kwargs.kwargs.get("json") or call_kwargs[1].get("json")
    user_msg = sent_payload["messages"][1]
    assert isinstance(user_msg["content"], list)
    assert user_msg["content"][1]["type"] == "image_url"
    assert "data:image/png;base64," in user_msg["content"][1]["image_url"]["url"]


def test_generate_content_too_large(client, mock_settings_attrs):
    """Content exceeding max size is rejected."""
    mock_settings_attrs["max_content_length"] = 10
    payload = {
        "content": "x" * 100,
        "content_type": "text/plain",
        "prompt": "Summarize.",
        "output_schema": {"summary": "str"},
    }

    with patch("llm_relay.main.settings") as mock_settings:
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is False
    assert "maximum size" in data["error"]


def test_generate_empty_document(client, mock_settings_attrs):
    """Empty document content returns an error, not hallucinated output."""
    import base64
    import io

    import docx as docx_lib

    doc = docx_lib.Document()  # empty doc, no paragraphs
    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    payload = {
        "content": b64,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "prompt": "Summarize.",
        "output_schema": {"summary": "str"},
    }

    with patch("llm_relay.main.settings") as mock_settings:
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is False
    assert "empty" in data["error"].lower()


def test_generate_docx_end_to_end(client, mock_settings_attrs):
    """Real .docx file through the full generate endpoint."""
    import base64
    import io

    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_paragraph("Zaaktype: bezwaarschrift tegen omgevingsvergunning.")
    doc.add_paragraph("Status: in behandeling bij de commissie.")
    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    payload = {
        "content": b64,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "prompt": "Extract the zaaktype and status.",
        "output_schema": {"zaaktype": "str", "status": "str"},
    }
    llm_data = {"zaaktype": "bezwaarschrift", "status": "in behandeling"}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response) as mock_post,
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert data["data"] == llm_data

    # Verify extracted text was sent to the LLM (not raw base64)
    sent = mock_post.call_args.kwargs.get("json") or mock_post.call_args[1].get("json")
    user_content = sent["messages"][1]["content"]
    assert "bezwaarschrift" in user_content
    assert "in behandeling" in user_content


def test_generate_xlsx_end_to_end(client, mock_settings_attrs):
    """Real .xlsx file through the full generate endpoint."""
    import base64
    import io

    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Metadata"
    ws["A1"] = "Kenmerk"
    ws["B1"] = "ZK-2026-001"
    ws["A2"] = "Beschrijving"
    ws["B2"] = "Aanvraag omgevingsvergunning"
    buf = io.BytesIO()
    wb.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    payload = {
        "content": b64,
        "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "prompt": "Extract the kenmerk and beschrijving.",
        "output_schema": {"kenmerk": "str", "beschrijving": "str"},
    }
    llm_data = {"kenmerk": "ZK-2026-001", "beschrijving": "Aanvraag omgevingsvergunning"}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response) as mock_post,
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True

    # Verify extracted text was sent (not raw base64)
    sent = mock_post.call_args.kwargs.get("json") or mock_post.call_args[1].get("json")
    user_content = sent["messages"][1]["content"]
    assert "ZK-2026-001" in user_content
    assert "omgevingsvergunning" in user_content


def test_generate_real_png_end_to_end(client, mock_settings_attrs):
    """Real minimal PNG through the full generate endpoint with vision."""
    import base64
    import struct
    import zlib

    # Build a valid 1x1 red PNG
    def _make_png():
        sig = b"\x89PNG\r\n\x1a\n"

        def chunk(ctype, data):
            c = ctype + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1 RGB
        raw_row = b"\x00\xff\x00\x00"  # filter byte + red pixel
        idat = zlib.compress(raw_row)
        return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")

    png_bytes = _make_png()
    b64 = base64.b64encode(png_bytes).decode()

    payload = {
        "content": b64,
        "content_type": "image/png",
        "prompt": "Describe the image.",
        "output_schema": {"description": "str", "color": "str"},
    }
    llm_data = {"description": "A single red pixel", "color": "red"}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response) as mock_post,
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    assert response.json()["success"] is True

    # Verify vision format: image_url with base64 data URI
    sent = mock_post.call_args.kwargs.get("json") or mock_post.call_args[1].get("json")
    user_msg = sent["messages"][1]
    assert isinstance(user_msg["content"], list)
    img_part = user_msg["content"][1]
    assert img_part["type"] == "image_url"
    assert img_part["image_url"]["url"].startswith("data:image/png;base64,")


def test_generate_eml_end_to_end(client, mock_settings_attrs):
    """Real .eml content through the full generate endpoint."""
    import base64

    eml = (
        "From: burger@example.nl\r\n"
        "To: gemeente@example.nl\r\n"
        "Subject: Bezwaar tegen besluit\r\n"
        "Content-Type: text/plain\r\n"
        "\r\n"
        "Hierbij maak ik bezwaar tegen het besluit van 15 maart 2026.\r\n"
    )
    b64 = base64.b64encode(eml.encode()).decode()

    payload = {
        "content": b64,
        "content_type": "message/rfc822",
        "attachment_type": "item",
        "prompt": "Extract sender, subject, and main request.",
        "output_schema": {"sender": "str", "subject": "str", "request": "str"},
    }
    llm_data = {"sender": "burger@example.nl", "subject": "Bezwaar tegen besluit", "request": "bezwaar"}
    mock_response = _mock_openrouter_response(llm_data)

    with (
        patch("llm_relay.main.settings") as mock_settings,
        patch("httpx.AsyncClient.post", new_callable=AsyncMock, return_value=mock_response) as mock_post,
    ):
        for k, v in mock_settings_attrs.items():
            setattr(mock_settings, k, v)
        response = client.post("/api/v1/generate", json=payload)

    assert response.status_code == 200
    assert response.json()["success"] is True

    # Verify extracted email text was sent, with email context hint
    sent = mock_post.call_args.kwargs.get("json") or mock_post.call_args[1].get("json")
    user_content = sent["messages"][1]["content"]
    assert "burger@example.nl" in user_content
    assert "Bezwaar" in user_content
    assert "email" in user_content.lower()  # context hint for attachment_type=item
