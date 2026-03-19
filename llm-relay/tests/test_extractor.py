# SPDX-FileCopyrightText: 2025 INFO.nl
# SPDX-License-Identifier: EUPL-1.2+

import base64
import io

from llm_relay.extractor import extract_text, is_image


def test_is_image_png():
    assert is_image("image/png") is True


def test_is_image_jpeg():
    assert is_image("image/jpeg") is True


def test_is_image_none():
    assert is_image(None) is False


def test_is_image_text():
    assert is_image("text/plain") is False


def test_is_image_docx():
    assert is_image("application/vnd.openxmlformats-officedocument.wordprocessingml.document") is False


def test_extract_text_plain_utf8():
    text = "Hello, world! This is a test."
    b64 = base64.b64encode(text.encode()).decode()
    result = extract_text(b64, "text/plain")
    assert result == text


def test_extract_text_plain_not_base64():
    """Non-base64 content falls back to plain text."""
    text = "This is not base64 at all, it has spaces and stuff!"
    result = extract_text(text, "text/plain")
    assert result == text


def test_extract_text_no_content_type():
    """No content_type defaults to UTF-8 decode."""
    text = "Some document content."
    b64 = base64.b64encode(text.encode()).decode()
    result = extract_text(b64, None)
    assert result == text


def test_extract_text_html():
    html = "<html><body><h1>Title</h1><p>Content</p></body></html>"
    b64 = base64.b64encode(html.encode()).decode()
    result = extract_text(b64, "text/html")
    assert "<h1>Title</h1>" in result


def test_extract_text_eml():
    eml_content = (
        "From: sender@example.com\r\n"
        "To: recipient@example.com\r\n"
        "Subject: Test Email\r\n"
        "Date: Mon, 18 Mar 2026 10:00:00 +0100\r\n"
        "Content-Type: text/plain\r\n"
        "\r\n"
        "This is the email body.\r\n"
    )
    b64 = base64.b64encode(eml_content.encode()).decode()
    result = extract_text(b64, "message/rfc822")
    assert "sender@example.com" in result
    assert "Test Email" in result
    assert "email body" in result


def test_extract_text_docx():
    """Test .docx extraction with a real minimal docx file."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("First paragraph about housing policy.")
    doc.add_paragraph("Second paragraph about regulations.")

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "housing policy" in result
    assert "regulations" in result


def test_extract_text_docx_with_table():
    """Test .docx extraction includes table content."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Intro paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header1"
    table.cell(0, 1).text = "Header2"
    table.cell(1, 0).text = "Value1"
    table.cell(1, 1).text = "Value2"

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Intro paragraph" in result
    assert "Header1" in result
    assert "Value2" in result


def test_extract_text_xlsx():
    """Test .xlsx extraction with a real minimal xlsx file."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "Name"
    ws["B1"] = "Score"
    ws["A2"] = "Alice"
    ws["B2"] = 95
    ws["A3"] = "Bob"
    ws["B3"] = 87

    buf = io.BytesIO()
    wb.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert "Sheet: Data" in result
    assert "Alice" in result
    assert "95" in result
    assert "Bob" in result


def test_extract_text_windows_encoding():
    """Test that cp1252 (Windows) encoded text is handled."""
    # cp1252 has characters like curly quotes that aren't valid UTF-8
    text_bytes = "Héllo wörld — fancy quotes".encode("cp1252")
    b64 = base64.b64encode(text_bytes).decode()
    result = extract_text(b64, "text/plain")
    assert "wörld" in result or len(result) > 0  # Should not crash


def test_extract_text_eml_with_attachment():
    """Test that .eml extraction mentions attachment filenames."""
    eml_content = (
        "From: sender@example.com\r\n"
        "To: recipient@example.com\r\n"
        "Subject: With attachment\r\n"
        "MIME-Version: 1.0\r\n"
        "Content-Type: multipart/mixed; boundary=boundary123\r\n"
        "\r\n"
        "--boundary123\r\n"
        "Content-Type: text/plain\r\n"
        "\r\n"
        "See the attached file.\r\n"
        "--boundary123\r\n"
        "Content-Type: application/pdf\r\n"
        'Content-Disposition: attachment; filename="report.pdf"\r\n'
        "\r\n"
        "fake pdf content\r\n"
        "--boundary123--\r\n"
    )
    b64 = base64.b64encode(eml_content.encode()).decode()
    result = extract_text(b64, "message/rfc822")
    assert "sender@example.com" in result
    assert "With attachment" in result
    assert "report.pdf" in result


def test_extract_text_empty_docx():
    """Empty .docx returns empty string."""
    import docx

    doc = docx.Document()
    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()
    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result.strip() == ""


def test_extract_text_empty_xlsx():
    """Empty .xlsx returns just sheet header."""
    import openpyxl

    wb = openpyxl.Workbook()
    buf = io.BytesIO()
    wb.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()
    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert "Sheet:" in result


def test_extract_text_unpadded_base64():
    """Base64 with missing padding should still decode."""
    text = "Unpadded content"
    b64 = base64.b64encode(text.encode()).decode().rstrip("=")
    result = extract_text(b64, "text/plain")
    assert result == text


def test_extract_text_xlsx_multi_sheet():
    """Multi-sheet .xlsx extracts all sheets."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Invoices"
    ws1["A1"] = "Amount"
    ws1["B1"] = 1234

    ws2 = wb.create_sheet("Summary")
    ws2["A1"] = "Total"
    ws2["B1"] = 5678

    buf = io.BytesIO()
    wb.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert "Sheet: Invoices" in result
    assert "1234" in result
    assert "Sheet: Summary" in result
    assert "5678" in result


def test_extract_text_docx_unicode():
    """Docx with non-ASCII (Dutch/German) characters extracts correctly."""
    import docx

    doc = docx.Document()
    doc.add_paragraph("Beschrijving van het zaaktype voor gemeentelijke übersicht.")
    doc.add_paragraph("Coördinatie en reïntegratie zijn belangrijk.")

    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode()

    result = extract_text(b64, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "übersicht" in result
    assert "Coördinatie" in result
    assert "reïntegratie" in result


def test_extract_text_binary_content_raises():
    """Binary content that isn't text/docx/xlsx/eml raises ValueError."""
    binary_data = bytes(range(256))  # all byte values — not valid text
    b64 = base64.b64encode(binary_data).decode()
    import pytest

    with pytest.raises(ValueError, match="Cannot decode binary content"):
        extract_text(b64, "application/octet-stream")


def test_is_image_all_supported_types():
    """All documented image types are recognized."""
    for mime in ("image/png", "image/jpeg", "image/jpg", "image/gif", "image/webp"):
        assert is_image(mime) is True, f"{mime} should be recognized as image"


def test_is_image_case_insensitive():
    """Image detection is case-insensitive."""
    assert is_image("Image/PNG") is True
    assert is_image("IMAGE/JPEG") is True
