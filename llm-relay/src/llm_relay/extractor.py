# SPDX-FileCopyrightText: 2025 INFO.nl
# SPDX-License-Identifier: EUPL-1.2+

import base64
import email
import io
import logging
from email import policy

logger = logging.getLogger(__name__)

# MIME types that are images (sent via vision API, not extracted)
IMAGE_MIME_TYPES = frozenset(
    {
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/gif",
        "image/webp",
    }
)

# MIME types for Office documents that need extraction
# Note: only OOXML (.docx/.xlsx) is supported — legacy .doc/.xls (application/msword,
# application/vnd.ms-excel) are binary formats that python-docx/openpyxl cannot read.
DOCX_MIME_TYPES = frozenset(
    {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
)

XLSX_MIME_TYPES = frozenset(
    {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
)

EML_MIME_TYPES = frozenset(
    {
        "message/rfc822",
    }
)


def is_image(content_type: str | None) -> bool:
    return content_type is not None and content_type.lower() in IMAGE_MIME_TYPES


def extract_text(raw_b64: str, content_type: str | None) -> str:
    """Extract readable text from base64-encoded content based on MIME type.

    For images, this should NOT be called — use is_image() to check first.
    """
    ct = (content_type or "").lower()

    # Try to decode base64 first
    try:
        raw_bytes = base64.b64decode(raw_b64, validate=True)
    except Exception:
        # Not base64 — treat as plain text
        logger.info("Content is not base64 — using as plain text")
        return raw_b64

    if ct in DOCX_MIME_TYPES:
        return _extract_docx(raw_bytes)
    elif ct in XLSX_MIME_TYPES:
        return _extract_xlsx(raw_bytes)
    elif ct in EML_MIME_TYPES:
        return _extract_eml(raw_bytes)
    else:
        # Default: try UTF-8 decode
        try:
            return raw_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Try common Windows encoding
            try:
                return raw_bytes.decode("cp1252")
            except UnicodeDecodeError:
                logger.warning("Could not decode content as text for content_type=%s", content_type)
                return raw_b64


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx file."""
    import docx

    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_xlsx(data: bytes) -> str:
    """Extract text from an .xlsx file as CSV-like representation."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_eml(data: bytes) -> str:
    """Extract text from an .eml email file."""
    msg = email.message_from_bytes(data, policy=policy.default)
    parts = []

    # Headers
    for header in ("From", "To", "Cc", "Subject", "Date"):
        value = msg.get(header)
        if value:
            parts.append(f"{header}: {value}")

    parts.append("")  # blank line after headers

    # Body
    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        content = body.get_content()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        parts.append(content)

    # Note any attachments
    attachments = [part.get_filename() for part in msg.iter_attachments() if part.get_filename()]
    if attachments:
        parts.append(f"\n[Attachments: {', '.join(attachments)}]")

    return "\n".join(parts)
